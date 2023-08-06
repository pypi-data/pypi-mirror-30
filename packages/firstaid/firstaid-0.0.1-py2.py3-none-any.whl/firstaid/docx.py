from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
import pandas as pd
from PIL import Image


class DOCX():
    def __init__(self):
        self.doc = Document()
        # page dimensions
        self.PAGE_WIDTH = self.doc.sections[0].page_width.pt
        self.PAGE_HEIGHT = self.doc.sections[0].page_height.pt
        # margins
        self.PAGE_LEFT_MARGIN = self.doc.sections[0].left_margin.pt
        self.PAGE_RIGHT_MARGIN = self.doc.sections[0].right_margin.pt
        self.PAGE_TOP_MARGIN = self.doc.sections[0].top_margin.pt
        self.PAGE_BOTTOM_MARGIN = self.doc.sections[0].bottom_margin.pt
        self.PAGE_HEADER_DISTANCE = self.doc.sections[0].header_distance.pt
        self.PAGE_FOOTER_DISTANCE = self.doc.sections[0].footer_distance.pt
        # dimensions including margins
        self.ACTUAL_PAGE_WIDTH = self.PAGE_WIDTH - self.PAGE_LEFT_MARGIN - self.PAGE_RIGHT_MARGIN
        self.ACTUAL_PAGE_HEIGHT = self.PAGE_HEIGHT - self.PAGE_TOP_MARGIN - self.PAGE_BOTTOM_MARGIN - \
                                  self.PAGE_HEADER_DISTANCE - self.PAGE_FOOTER_DISTANCE
        # style
        self.TABLE_STYLE = 'LightShading-Accent1'
        self.FIRST_LINE_INDENT = Cm(1)

    def add_paragraph(self, content):
        """
        add paragraphs to page
        :param content: list of strings. ex) ['first paragraph', 'second paragraph']
        """
        for con in content:
            paragraph = self.doc.add_paragraph(con)
            if len(con) > 45:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = self.FIRST_LINE_INDENT

    def add_title(self, content):
        """
        add heading to page
        :param content: string. ex) 'Heading 1'
        """
        self.doc.add_heading(content)

    def add_image(self, filename):
        """
        add image to page
        :param filename: file name
        """
        im = Image.open(filename)
        sd = {'width': self.ACTUAL_PAGE_WIDTH, 'height': self.ACTUAL_PAGE_HEIGHT}
        d = {'width': im.width, 'height': im.height}

        max_key = max(d.keys(), key=lambda k: d[k])
        all_keys = list(d.keys())
        all_keys.pop(all_keys.index(max_key))
        counterpart = all_keys[0]

        if d[max_key] > sd[max_key]:  # larger than suggested
            new_max_value = sd[max_key]
        else:  # smaller than suggested
            new_max_value = d[max_key]

        d['new_{}'.format(counterpart)] = new_max_value * d[counterpart] / d[max_key]
        d['new_{}'.format(max_key)] = new_max_value

        fig = BytesIO()
        im.save(fig, 'PNG')
        fig.seek(0)
        self.doc.add_picture(fig, Pt(d['new_width']), Pt(d['new_height']))
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # align image to center of page

    def add_table(self, df):
        """
        add table to slide
        :param df: pandas DataFrame
        """
        table = self.doc.add_table(df.shape[0] + 1, df.shape[1] + 1)
        table.style = self.TABLE_STYLE
        # column title
        for i, col in enumerate(df.columns):
            table.cell(0, i + 1).text = str(col)
        # row title
        for i, row in enumerate(df.index):
            table.cell(i + 1, 0).text = str(row)
        # values
        for i, row in enumerate(df.index):
            for e, item in enumerate(df.loc[row]):
                table.cell(i + 1, e + 1).text = str(item)

        return table

    def page_break(self):
        """
        add page break
        """
        self.doc.add_page_break()

    def blank_line(self):
        """
        add new line
        """
        self.doc.add_paragraph('')

    def add(self, content='', image=''):
        """
        main function to add content to page
        :param content: content to insert in page
        :param image: only used if content is image files
        """
        # image
        if image:
            self.add_image(image)
        # string
        if isinstance(content, list):
            self.add_paragraph(content)
        # title
        elif isinstance(content, str):
            self.add_title(content)
        # table
        elif isinstance(content, pd.DataFrame):
            self.add_table(content)

    def save(self, file_name):
        """
        save document
        :param file_name: file name. ex) 'test.docx'
        """
        self.doc.save(file_name)
