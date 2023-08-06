#!/usr/bin/env python3
import os, sys, logging
from docx import Document
from docx.shared import Inches
import imghdr, struct, math
from .LogGen import set_up_logging
from .XtraFunctions import GetDate, OrderMe, cli_progress_test

logger = logging.getLogger(__name__)

class pics2word:

    def __init__(self):
        # set default values on instantiation, until changed with CL args
        logger.debug("Running constructor functions")
        self.SetPath()
        self.SetTitle() 
        self.SetPicWidth()
        self.SetPicHeight()
        self.SetFormat()
        self.SetTableWidth()

    def SetPath(self, Path=os.getcwd()):
        # Default path is the current working directory on the command line
        logger.info("Getting pictures from %s." % Path)
        self.path = os.path.abspath(Path)
        # Once we have the path, we get the pictures in that path
        self.GetPicsInPath()

    def SetTitle(self, title="PhotoDoc", date='y'):
        # if date begins with 'y', append title with date
        if date[0] == 'y':
            Today = GetDate()
            self.title = title + "_" + str(Today)
        else:
            self.title = title

    def SetPicWidth(self,Width=4):
        # TODO set a default!
        self.width = Width

    def SetPicHeight(self,Height=4):
        # TODO set a default!
        self.height = Height

    def SetTableWidth(self, Columns=2):
        # TODO set a default!
        self.tablecolumns = Columns
    
    def SetFormat(self, format="normal"):
        logger.debug("Setting format to %s." % format)
        if format[0].lower() == 't':
            self.format = "table"
        elif format[0].lower() == 'n':
            self.format = "normal"
        else:
            raise ValueError("Please enter a valid format for '-f' i.e. \"Normal\" or \"Table\"")

    def WriteDoc(self):
        if self.format[0].lower() == 't':
            logger.info("Writing to table.")
            self.WriteTable()
        else:
            logger.info("Writing to document normally.")
            self.WriteNormal()

    def WriteNormal(self):
        logger.info("Creating word document.")
        document = Document()
        p = document.add_paragraph()
        Path = self.path
        # Todo check if numbered and sort appropriately
        logger.debug("Sorting list of %s pictures." % len(self.pics))
        PicList = OrderMe(self.pics) # Sort pics into an order
        i=0
        for Pic in PicList:
            logger.debug("Writing picture: %s." % Pic)
            FullImageandPath = os.path.join(Path,Pic)
            r = p.add_run()
            logger.debug("Checking if %s is portrait." % Pic)
            isPortrait = self.IsPortrait(FullImageandPath)
            logger.debug("Adding %s to file." % Pic)
            if isPortrait:
                r.add_picture(FullImageandPath,height=Inches(self.height))
            else:
                r.add_picture(FullImageandPath,width=Inches(self.width))
            logger.debug("Adding description: %s." % Pic.split('.')[0])
            p.add_run("\n"+Pic.split('.')[0]+"\n")
            # update progress
            logger.debug("writing loading bar picture %s of %s." % (i, len(PicList)))
            cli_progress_test(cur_val=i,end_val=len(PicList),suffix=("Writing pic %s of %s." % (i, len(PicList))))
            i += 1
        logger.info("Saving document as %s.docx" % self.title)
        document.save(self.title + '.docx')

    def WriteTable(self):
        logger.info("Creating word document.")
        document = Document()
        Path = self.path
        logger.debug("Sorting list of %s pictures." % len(self.pics))
        PicList = OrderMe(self.pics) # Sort pics into an order
        logger.info("Calculating number of rows.")
        numRows = self.GetNumberofRows()
        logger.info("Adding table of %s rows and %s columns." % (numRows, self.tablecolumns))
        table = document.add_table(rows=numRows, cols=self.tablecolumns)
        i=0
        Row_Index = 0 # Resets every iteration
        # list[start:stop:step] pastes the picture in every 2nd cell
        for row in table.rows[::2]:
            Col_Index = 0 # Resets every iteration
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    try:
                        Pic = PicList[i]
                        logger.debug("Writing %s in row %s cell %s paragraph %s." % (Pic, row, cell, paragraph))
                        FullImageandPath = os.path.join(str(Path),str(Pic))
                        r = paragraph.add_run()

                        logger.debug("Checking if %s is portrait." % Pic)
                        isPortrait = self.IsPortrait(FullImageandPath)
                        logger.debug("Adding %s to file." % Pic)
                        if isPortrait:
                            r.add_picture(FullImageandPath,height=Inches(self.height))
                        else:
                            r.add_picture(FullImageandPath,width=Inches(self.width))

                        # Offset 1 row down and add description
                        logger.debug("Writing description %s in row %s, column %s" % (Pic.split('.')[0], row, cell))
                        table.cell(row_idx=Row_Index + 1,col_idx=Col_Index).text = Pic.split('.')[0]
                        # Update user of progress
                        logger.debug("writing loading bar picture %s of %s." % (i, len(PicList)))
                        cli_progress_test(cur_val=i,end_val=len(PicList),suffix=("Writing pic %s of %s." % (i, len(PicList))))
                    except IndexError:
                        # we incur an index error at the end of the picture list
                        # hence, we will simply pass and do nothing with the remaining empty cells
                        logger.warning("Index Error on picture %s indicating that there are remaining cells but no new pictures." % Pic)
                        pass
                    except:
                        print("Error in copying picture to table: %s" % sys.exc_info()[0])
                        logger.error("Error in copying picture to table: %s" % sys.exc_info()[0])
                        raise
                Col_Index += 1
                i += 1
            Row_Index += 2
        logger.info("Saving document as %s.docx" % self.title)
        document.save(self.title + '.docx')
    
    def GetPicsInPath(self):
        # Runs from the set path function
        self.pics = []
        ValidExtList = [".jpg",".jpeg",".png",".bmp",".gif",".JPG",".JPEG",".PNG",".BMP",".GIF"]
        for file in os.listdir(self.path):
            for ValidExt in ValidExtList:
                if file.endswith(ValidExt):
                    logger.debug("Adding %s to PicList" % file)
                    self.pics.append(file)
        return self.pics

    def IsPortrait(self, fname):
        """Determine the image type of fhandle and return its size."""
        with open(fname, 'rb') as fhandle:
            head = fhandle.read(24)
            if len(head) != 24:
                return
            if imghdr.what(fname) == 'png':
                check = struct.unpack('>i', head[4:8])[0]
                if check != 0x0d0a1a0a:
                    return
                width, height = struct.unpack('>ii', head[16:24])
            elif imghdr.what(fname) == 'gif':
                width, height = struct.unpack('<HH', head[6:10])
            elif imghdr.what(fname) == 'jpeg':
                try:
                    fhandle.seek(0) # Read 0xff next
                    size = 2
                    ftype = 0
                    while not 0xc0 <= ftype <= 0xcf:
                        fhandle.seek(size, 1)
                        byte = fhandle.read(1)
                        while ord(byte) == 0xff:
                            byte = fhandle.read(1)
                        ftype = ord(byte)
                        size = struct.unpack('>H', fhandle.read(2))[0] - 2
                    # We are at a SOFn block
                    fhandle.seek(1, 1)  # Skip `precision' byte.
                    height, width = struct.unpack('>HH', fhandle.read(4))
                except Exception: #IGNORE:W0703
                    return
            else:
                return
            if width/height > 1:
                return False
            else:
                return True  

    def GetNumberofRows(self):
        cols = self.tablecolumns
        NumofPics = len(self.pics)
        return int(math.ceil(NumofPics / cols)) * 2
