from docx import Document
from docx.shared import Inches
import docx.shared
from datetime import datetime
import os

now = datetime.today()
dates = now.strftime('%A, %d %b %Y')
times = now.strftime("%I: %M: %S %p")
date_time = dates + " " + times

class MyDocument:
    """ Main class, creates an instance of docx.Document,
    and instantiates the current date and time in strftime formats,
    '%A, %d %b %Y' and '%I: %M: %S %p' respectively)
    """
    def __init__(self, date_time=date_time):
        ''' Constructor for MyDocument'''

        self.doc = Document("default.docx")
        self.date_time = date_time

    def add_title(self,title):
        ''' Adds a heading with level= 0, which represents a titile of the document'''
        self.title = title
        return self.doc.add_heading(self.title, level = 1)
    def add_heading(self, text, level):
        return self.doc.add_heading(text, level=level)

    def add_date_time(self, date_time = None):
        """ Args: date_time, default= None. Returns a date object in a paragraph"""
        if date_time:
            return self.doc.add_paragraph(date_time)

        return self.doc.add_paragraph(self.date_time)

    def add_table(self, table_heading, headers, records):
        """ Returns an Ms-Word table, and inserts it to the document
        Args: table_heading, headers(should be a list), records(should be an
        iterable ie a list of tuples or tuple of tuples)"""
        self.doc.add_heading(table_heading, level = 3)

        table = self.doc.add_table(rows=1, cols= len(headers))

        hdr_cells = table.rows[0].cells

        for index, header in enumerate(headers):
            hdr_cells[index].text = str(header)

        for item in records:
            row_cells = table.add_row().cells
            for i in range(len(item)):
                row_cells[i].text = str(item[i])

        return table


    def add_image(self, image_path, width=None, height=None):
        """ Adds an image object to the word document, specified by the absolute
        path to the file.
        Args: Specified in inches as *args(width, height) or as **kwargs
        width=x, height =y
        """
        if width !=None and height !=None:
            return self.doc.add_picture(os.path.abspath(image_path), width=Inches(width),
                height = Inches(height))
        else:
            return self.doc.add_picture(os.path.abspath(image_path))

    def add_text(self, paragraph,text):
        """Args: paragraph, text;
        Returns new text added to the paragraph specified.
        New line(\n), tabs(\t) and carriage return(\r) characters
        have to be explicitly passed.
        """
        return paragraph.add_run(text)

    def add_paragraph(self, text):
        return self.doc.add_paragraph(text)

    def save(self, filename):
        """ Saves the file by the absolute filename specified as the only argument"""
        try:
            self.doc.save(os.path.abspath(filename))
        except PermissionError:
            print("File Permision Error...Close file first")

    def __str__(self):
        """ Dander method ---
        Returns the class documentation"""
        return str(help(MyDocument))

if __name__ == '__main__':
    doc = MyDocument()
    doc.add_title("Add Title")
    doc.save("Demo.docx")
    print(doc)

