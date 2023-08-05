from azure.common.credentials import ServicePrincipalCredentials
from azure.mgmt.resource import ResourceManagementClient
from docx import Document
from azure.mgmt.compute import ComputeManagementClient
from azure.mgmt.network import NetworkManagementClient
import yaml
import os


# Confguation section
Config = yaml.load(open('config.yml'))
SUBID = Config['odit']['provider']['azure']['auth']['account']
TENANT_ID = Config['odit']['provider']['azure']['auth']['other']
CLIENT = Config['odit']['provider']['azure']['auth']['secret']
KEY = Config['odit']['provider']['azure']['auth']['key']



subscription_id = os.environ.get(
    'AZURE_SUBSCRIPTION_ID',
    SUBID)
credentials = ServicePrincipalCredentials(
    client_id=CLIENT,
    secret=KEY,
    tenant=TENANT_ID
)




def set_subscription():
    return subscription_id


def set_tenant():
    return TENANT_ID


def resource_groups():
    client = ResourceManagementClient(credentials, subscription_id)
    data = client.resource_groups.list()
    return data


def resources(name):
    rclient = ResourceManagementClient(credentials, subscription_id)
    rdata = rclient.resources.list_by_resource_group(name)
    return rdata


def get_vnet():
    network_client = NetworkManagementClient(credentials, subscription_id)
    data = network_client.virtual_networks.list_all()
    return data

def get_vms():
    compute_client = ComputeManagementClient(credentials, subscription_id)
    data = compute_client.virtual_machines.list_all()
    return data


def generate_design_document():
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    # Document heading
    document.add_heading('Build Document', 0)
    # Document Owner
    p = document.add_paragraph('This design document is created by Odit "The Nocturnal" ')
    # Account Section
    document.add_heading('Account Details', level=1)
    # Subscription Section
    document.add_heading('Subscription', level=2)
    SubData = set_subscription()
    SD = document.add_paragraph(SubData)
    SD.add_run('bold').bold = True
    # Tenant Section
    document.add_heading('Tenant ID', level=2)
    TenData = set_tenant()
    TN = document.add_paragraph(TenData)
    TN.add_run('bold').bold = True
    # Resource Groups
    document.add_heading('Resource Groups', level=1)
    RGData = resource_groups()
    x = 0

    for r in RGData:
        document.add_heading(r.name, level=2)
        document.add_heading("Overview", level=3)
        overview_statement = '%s is located at %s region & is managed by %s. ' % (r.name, r.location, r.managed_by)
        document.add_paragraph(overview_statement)
        state_statement = 'Last provisioning state for %s with id %s was %s' % (
            r.name, r.id, r.properties.provisioning_state)
        document.add_paragraph(state_statement)
        d = resources(r.name)
        y = 0
        for e in d:
            y = y + 1
        count_statement = '%s contains %s resources' % (r.name, y)
        document.add_paragraph(count_statement)
        m = resources(r.name)
        document.add_heading("List of resources", level=4)
        for i in m:
            detail_statement = i.name
            document.add_paragraph(detail_statement, style='ListBullet')
        document.add_heading("Resources Classification", level=3)
        q = resources(r.name)
        for z in q:
            classification = '%s is a %s, kind of resource is %s. ' % (z.name, z.type, z.kind)
            document.add_heading("SKU", level=4)

            if z.sku == None:
                sku = "There is no Sku associated by Microsoft for this resource."
                document.add_paragraph(sku)
            else:
                sku = 'Sku details for this resource are as followed:-'
                document.add_paragraph(sku)
                document.add_paragraph('Capacity: %s' %z.sku.capacity, style='ListBullet')
                document.add_paragraph('Name: %s' %z.sku.name, style='ListBullet')
                document.add_paragraph('Family: %s' % z.sku.family, style='ListBullet')
                document.add_paragraph('Tier: %s' % z.sku.tier, style='ListBullet')
                document.add_paragraph('Model: %s' % z.sku.model, style='ListBullet')
                document.add_paragraph('Size: %s' % z.sku.size, style='ListBullet')

            document.add_heading("Identifier", level=4)
            idt = 'identifier for this resource is %s' % z.id
            document.add_paragraph(idt)
            document.add_heading("Plan", level=4)
            if z.plan == None:
                plan = "There is no plan associated for this resource."
                document.add_paragraph(plan)
            else:
                plan = 'Plan details for this resource are as followed:-'
                document.add_paragraph(plan)
                document.add_paragraph('Publisher: %s' % z.plan.publisher, style='ListBullet')
                document.add_paragraph('Product: %s' % z.plan.product, style='ListBullet')
                document.add_paragraph('Promotion Code: %s' % z.plan.promotion_code, style='ListBullet')
                document.add_paragraph('Name: %s' % z.plan.name, style='ListBullet')
            document.add_heading("Management", level=4)
            mgngby = 'This resource is managed by %s' % z.managed_by
            document.add_paragraph(mgngby)
            document.add_heading("Identity", level=4)
            idnt = 'Identity in use is %s' %z.identity
            document.add_paragraph(idnt)

    # Virtual Networks
    document.add_heading('Networks', level=1)
    net = get_vnet()
    for n in net:

        document.add_heading(n.name, level=2)
        document.add_heading("Address Space", level=3)
        addr = 'Address Space for this Network is %s' % n.address_space.address_prefixes[0]
        document.add_paragraph(addr)
        document.add_heading("Subnets", level=3)
        for sub in n.subnets:
            document.add_heading(sub.name, level=4)
            document.add_paragraph('Address Prefix: %s' % sub.address_prefix)
            document.add_paragraph('Provisioning State: %s' % sub.provisioning_state)
            document.add_paragraph('etag: %s' % sub.etag)
            document.add_paragraph('Route table: %s' %sub.route_table)
            document.add_paragraph('Service Endpoints: %s' % sub.service_endpoints)
            document.add_paragraph('Resource Id: %s' % sub.id)
            document.add_paragraph('Network Security Group: %s' % sub.network_security_group)
            if sub.ip_configurations == None:
                document.add_heading("Ip Configuration", level=5)
                document.add_paragraph("There is no network specific Ip configuration for this subnet.")
            else:
                for ip in sub.ip_configurations:
                    document.add_heading("Ip Configuration", level=5)
                    document.add_paragraph('Subnet Id: %s' % ip.id)
                    document.add_paragraph('name: %s' % ip.name)
                    document.add_paragraph('Subnet: %s' % ip.subnet)
                    document.add_paragraph('Public Ip Address: %s' % ip.public_ip_address)
                    document.add_paragraph('Provisioning State: %s' % ip.provisioning_state)
                    document.add_paragraph('Private IP allocation Method: %s' %ip.private_ip_allocation_method)
                    document.add_paragraph('Private IP allocation Method: %s' % ip.private_ip_allocation_method)
                    document.add_paragraph('etag: %s' % ip.etag)
                    document.add_paragraph('Private Ip Address: %s' % ip.private_ip_address)
    # Virtual Machines
    document.add_heading('Virtual Machines', level=1)
    vms = get_vms()
    for vm in vms:
        document.add_heading(vm.name, level=2)
        document.add_paragraph('VM ID: %s' %vm.vm_id)
        aset = vm.availability_set
        if aset == None:
            document.add_paragraph("This vm is not part of availability set")
        else:
            document.add_paragraph("This machine is part of availability set")
            document.add_paragraph('Availability Set Id: %s' %vm.availability_set.id)
        document.add_heading("OS Details", level=3)
        document.add_paragraph('Admin Username: %s' % vm.os_profile.admin_username)
        document.add_paragraph('Admin password: %s' %vm.os_profile.admin_password)
        document.add_paragraph('Computer/Host Name: %s' % vm.os_profile.computer_name)
        document.add_paragraph('Custom date: %s' % vm.os_profile.custom_data)
        document.add_paragraph('Secrets: %s' % vm.os_profile.secrets)
        document.add_heading("Windows Configuration", level=4)
        if vm.os_profile.windows_configuration == None:
            document.add_paragraph("No Windows Configuration available on this virtual machine.")
        else:
            document.add_paragraph('WIN RM: %s' % vm.os_profile.windows_configuration.win_rm)
            document.add_paragraph('Additional unattended content: %s' % vm.os_profile.windows_configuration.additional_unattend_content)
            document.add_paragraph('Automatic updates enabled: %s' % vm.os_profile.windows_configuration.enable_automatic_updates)
            document.add_paragraph('Azure provisioning VM Agent enabled: %s' % vm.os_profile.windows_configuration.provision_vm_agent)
            document.add_paragraph('Time Zone: %s' % vm.os_profile.windows_configuration.time_zone)
        document.add_heading("Linux Configuration", level=4)
        if vm.os_profile.linux_configuration == None:
            document.add_paragraph("No Linux Configuration available on this virtual machine.")
        else:
            document.add_paragraph('Disable Password Authentication: %s' % vm.os_profile.linux_configuration.disable_password_authentication)
            for key in vm.os_profile.linux_configuration.ssh.public_keys:
                document.add_paragraph('SSH Public Key data: %s' %key.key_data)
                document.add_paragraph('SSH Public Path: %s' % key.path)
        document.add_heading("Storage profile", level=3)
        if vm.storage_profile == None:
            document.add_paragraph("No storage profile available on this virtual machine.")
        else:
            document.add_heading("OS Disk", level=4)
            document.add_paragraph('Disk name: %s ' %vm.storage_profile.os_disk.name)
            document.add_paragraph('Disk size (GB): %s ' % vm.storage_profile.os_disk.disk_size_gb)
            document.add_paragraph('Image: %s' %vm.storage_profile.os_disk.image)
            document.add_paragraph('Caching Type: %s' %vm.storage_profile.os_disk.caching)
            document.add_paragraph('VHD: %s' %vm.storage_profile.os_disk.vhd)
            document.add_paragraph('OS: %s' % vm.storage_profile.os_disk.os_type)
            document.add_paragraph('Create Option: %s' %vm.storage_profile.os_disk.create_option)
            document.add_paragraph('Encryption Settings: %s' % vm.storage_profile.os_disk.encryption_settings)
            document.add_paragraph('Managed Disk: %s' % vm.storage_profile.os_disk.managed_disk)
            document.add_heading("Image Reference", level=4)
            document.add_paragraph('SKU: %s' %vm.storage_profile.image_reference.sku)
            document.add_paragraph('Publisher: %s' % vm.storage_profile.image_reference.publisher)
            document.add_paragraph('Version: %s' % vm.storage_profile.image_reference.version)
            document.add_paragraph('Id: %s' % vm.storage_profile.image_reference.id)
            document.add_paragraph('Offer: %s' % vm.storage_profile.image_reference.offer)
            document.add_heading("Data Disks", level=4)
            f = vm.storage_profile.data_disks
            if not f:
                document.add_paragraph("No Data disks attached.")
            else:
                for disk in f:
                    document.add_paragraph('Managed Disk: %s' %disk.managed_disk)
                    document.add_paragraph('Disk Name: %s' % disk.name)
                    document.add_paragraph('Disk Image: %s' % disk.image)
                    document.add_paragraph('Caching: %s' % disk.caching)
                    document.add_paragraph('VHD: %s' % disk.vhd)
                    document.add_paragraph('Create option: %s' % disk.create_option)
            document.add_heading("Tags", level=3)
            if vm.tags == None:
                document.add_paragraph("There are no tags for this virtual machine")
            else:
                document.add_paragraph(vm.tags)
            document.add_heading("Diagnostic profile", level=3 )
            vmdp = vm.diagnostics_profile
            if not vmdp:
                document.add_paragraph("There are no diagnostic profile associated with this virtual machine")
            else:
                document.add_paragraph('Boot Diagnostics enabled: %s' % vm.diagnostics_profile.boot_diagnostics.enabled)
                document.add_paragraph('Boot Diagnostics enabled: %s' % vm.diagnostics_profile.boot_diagnostics.storage_uri)
            document.add_heading("Hardware profile", level=3)
            document.add_paragraph('VM Size: %s' %vm.hardware_profile.vm_size)
            document.add_heading("Provisioning State", level=3)
            document.add_paragraph('Last Provisioning State: %s' % vm.provisioning_state)
            document.add_heading("Zones", level=3)
            document.add_paragraph('Zone: %s' % vm.zones)
            document.add_heading("Network Profile", level=3)
            for inet in vm.network_profile.network_interfaces:
                document.add_paragraph('Network Interface Id: %s' %inet.id)
                document.add_paragraph('Primary: %s' % inet.primary)
            document.add_heading("Plan", level=3)
            vmplan = vm.plan
            if vmplan == None:
                document.add_paragraph("No plan associated with this VM")
            else:
                document.add_paragraph('Publisher: %s' %vmplan.publisher)
                document.add_paragraph('Product: %s' % vmplan.product)
                document.add_paragraph('Promotion code: %s' % vmplan.promotion_code)
                document.add_paragraph('Name: %s' % vmplan.name)
            document.add_heading("License Type", level=3)
            if vm.license_type == None:
                document.add_paragraph("No license type associated with this VM")
            else:
                document.add_paragraph('License type: %s' %vm.license_type)
            document.add_heading("Instance View", level=3)
            if vm.instance_view == None:
                document.add_paragraph("No instance view associated with this VM")
            else:
                document.add_paragraph('Instance View: %s' %vm.instance_view)
            document.add_heading("Type", level=3)
            if vm.type == None:
                document.add_paragraph("No vm type found")
            else:
                document.add_paragraph('VM Type: %s' % vm.type)
            document.add_heading("Location", level=3)
            if vm.location == None:
                document.add_paragraph("Unable to find a location")
            else:
                document.add_paragraph('VM Location: %s' % vm.location)



    document.save('odit_build_document.docx')



