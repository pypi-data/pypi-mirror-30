#!/bin/env python3
"""
Module for Lab Speleman reporting
"""
import matplotlib.pyplot as plt, re
from matplotlib.figure import Figure
import pandas as pd, re, inspect
from itertools import count
from collections import OrderedDict
from unittest.mock import Mock

# Settings
from .config import config
reportsDir = config['leopard']['reportdir']
csvsep = config['leopard']['csvsep']
csvdec = config['leopard']['csvdec']

class Section:
    """
    Report sections.
    Defines methods dealing with the structure of
    sections, and section specific output.

    tablehead => set to integer to only show DataFrame.head(tablehead) rows
    in this section

    For adding subsections, a method is provided.
    """
    def __init__(self,title,text,
                 figures=None,tables=None,subsections=None,code=None,
                 tablehead=None,tablecolumns=None,clearpage=False):
        self.title = title.strip()
        self.p = text.strip()
        self.figs = OrderedDict(figures) if figures else OrderedDict()
        self.tabs = OrderedDict(tables) if tables else OrderedDict()
        self.subs = subsections if subsections else []
        self.code = code
        self.settings = {'tablehead':tablehead,
                         'tablecolumns':tablecolumns,
                         'clearpage':clearpage,
                         'doubleslashnewline':False}
        # Initially set section upstream references to self
        self._parentSection = self
        self._reportSection = self

    def __repr__(self):
        return "<{} @ {}{}>".format(
            'Report' if 'sections' in dir(self) else 'Section',
            self.title[:50],
            '' if len(self.title)<=50 else '...'
        )
    
    def __getitem__(self,key):
        try: return self.subs[key]
        except TypeError:
            try: return self._subs[key]
            except (AttributeError,KeyError) as e:
                self._subs = {s.title:s for s in self.subs}
                return self._subs[key]

    def append(self,*args,toSection=None,**kwargs):
        """
        If toSection is None, section is appended to the main section/subs list.
        Else if toSection is int or (int,int,...), it gets added to the subs (subsection)
        list of the specified section.

        *args and **kwargs are processed by Section class initiation
        """
        if not toSection and toSection is not 0:
            s = Section(*args,**kwargs)
            self.subs.append(s)
            self.lastSection = s
            s._parentSection = self
            s._reportSection = self._reportSection
        else:
            if type(toSection) is int: toSection = (toSection,)
            s = self.subs[toSection[0]].append(*args,toSection=toSection[1:],**kwargs)
        return s

    @staticmethod
    def sectionWalker(section,callback,*args,walkTrace=tuple(),**kwargs):
        """
        callback needs to be a function that handles different 
        Section elements appropriately
        walkTrace needs to be a tuple, indicate the route to the section, e.g. (1,2,0)
        """
        callback(section,*args,walkTrace=walkTrace,case='sectionmain',**kwargs)
        c = count(1)
        for f in section.figs.items():
            callback(section,*args,walkTrace=walkTrace,case='figure',element=f,**kwargs)
        c = count(1)
        for t in section.tabs.items():
            callback(section,*args,walkTrace=walkTrace,case='table',element=t,**kwargs)
        c = count(1)
        for s in section.subs:
            Section.sectionWalker(s,callback,*args,walkTrace=walkTrace+(next(c),),**kwargs)

    def walkerWrapper(callback):
        def wrapper(*args,**kwargs):
            #args[0] => has to be the current walked section
            return Section.sectionWalker(args[0],callback,*args[1:],**kwargs)
        return wrapper

    @walkerWrapper
    def list(self,walkTrace=tuple(),case=None,element=None):
        if case == 'sectionmain': print(walkTrace,self.title)

    @walkerWrapper
    def listFigures(self,walkTrace=tuple(),case=None,element=None):
        if case == 'sectionmain': print(walkTrace,self.title)
        if case == 'figure':
            caption,fig = element
            try:
                print(walkTrace,fig._leopardref,caption)
            except AttributeError:
                fig._leopardref = next(self._reportSection._fignr)
                print(walkTrace,fig._leopardref,caption)

    @walkerWrapper
    def listTables(self,walkTrace=tuple(),case=None,element=None):
        if case == 'sectionmain': print(walkTrace,self.title)
        if case == 'table':
            caption,tab = element
            try:
                print(walkTrace,tab._leopardref,caption)
            except AttributeError:
                tab._leopardref = next(self._reportSection._tabnr)
                print(walkTrace,tab._leopardref,caption)

    def sectionOutZip(self,zipcontainer,zipdir='',figtype='png'):
        from io import StringIO
        with zipcontainer.open(zipdir+'section.txt',mode='w') as zipf:
            text = self.p if not self.settings['doubleslashnewline'] else self.p.replace('//','\n')
            zipf.write('# {}\n{}'.format(self.title,text).encode())
        c = count(1)
        for ftitle,f in self.figs.items():
            with zipcontainer.open(zipdir+'fig{}_{}.{}'.format(next(c),ftitle.replace(' ','_'),figtype),mode='w') as zipf:
                f.savefig(zipf,format=figtype,transparent=True)
        c = count(1)
        for ttitle,t in self.tabs.items():
            with zipcontainer.open(zipdir+'table{}_{}.csv'.format(next(c),ttitle.replace(' ','_')),mode='w') as zipf:
                b = StringIO()
                t.to_csv(b,sep=csvsep,decimal=csvdec)
                b.seek(0)
                zipf.write(b.read().encode())
        c = count(1)
        for s in self.subs:
            s.sectionOutZip(zipcontainer,'{}s{}_{}/'.format(zipdir,next(c),s.title.replace(' ','_')),figtype=figtype)

    @walkerWrapper
    def sectionsPDF(self,walkTrace=tuple(),case=None,element=None,doc=None):
        import pylatex as pl
        if case == 'sectionmain':
            if self.settings['clearpage']: doc.append(pl.utils.NoEscape(r'\clearpage'))
            with doc.create(pl.Section(self.title) if len(walkTrace) == 1 else
                            pl.Subsection(self.title) if len(walkTrace) == 2 else
                            pl.Subsubsection(self.title)):
                text = (self.p.replace('\n',' ').replace('//','\n')
                        if self.settings['doubleslashnewline'] else
                        renewliner(self.p))
                if r'\ref' not in text: doc.append(text)
                else:
                    figrefs = re.compile(r'\\ref\{figref\d+\}')
                    #latexcode = re.compile(r'&@\\.+')
                    lastpos = 0
                    for fr in figrefs.finditer(text):
                        doc.append(text[lastpos:fr.start()])
                        doc.append(pl.utils.NoEscape(text[fr.start():fr.end()]))
                        lastpos = fr.end()
                    doc.append(text[lastpos:])
                
        if case == 'figure':
            width = r'1\textwidth'
            figtitle,fig = element
            #if fig._suptitle: fig.suptitle('Figure {}: {}'.format(fig.number,fig._suptitle.get_text()))
            #figtitle = fig._suptitle.get_text() if fig._suptitle else ''
            #fig.suptitle('')
            with doc.create(pl.Figure(position='htbp')) as plot:
                plt.figure(fig.number)
                plot.add_plot(width=pl.NoEscape(width))
                plot.add_caption(figtitle)
                plot.append(pl.utils.NoEscape(r'\label{figref'+str(fig.number)+r'}'))
            #fig.suptitle(figtitle if figtitle else None)
            
        if case == 'table':
            caption,t = element
            t = pdSeriesToFrame(t) if type(t) == pd.Series else t
            if self.settings['tablehead']:
                t = t.head(self.settings['tablehead'])
            if self.settings['tablecolumns']:
                t = t[self.settings['tablecolumns']]
            with doc.create(pl.Table(position='ht')) as tablenv:
                tablenv.add_caption(caption)
                with doc.create(pl.Tabular('r|'+'l'*len(t.columns))) as table:
                    table.add_hline()
                    table.add_row(('',)+tuple(t.columns))
                    for row in t.to_records():
                        table.add_row(row)
                    table.add_hline(1)
                    #table.add_empty_row()

    @walkerWrapper
    def sectionsWord(self,walkTrace=tuple(),case=None,element=None,doc=None):
        from docx.shared import Inches
        from io import BytesIO
        #p.add_run('italic.').italic = True
                
        if case == 'sectionmain':
            if self.settings['clearpage']: doc.add_page_break()
            
            doc.add_heading(self.title, level = len(walkTrace))
            for p in renewliner(self.p).split('\n'):
                doc.add_paragraph(p)
                
        if case == 'figure':
            bf=BytesIO()
            figtitle,fig = element
            width = fig.get_size_inches()[0]
            width = Inches(width if width < 6 else 6)
            fig.savefig(bf)
            doc.add_picture(bf, width=width)
            doc.add_heading('Figure {}: {}'.format(
                fig._leopardref,
                figtitle),level=6)
            
        if case == 'table':
            caption,t = element
            tableref = t._leopardref
            t = pdSeriesToFrame(t) if type(t) == pd.Series else t
            if self.settings['tablehead']:
                t = t.head(self.settings['tablehead'])
            if self.settings['tablecolumns']:
                t = t[self.settings['tablecolumns']]

            doc.add_heading('Table {}: {}'.format(
                tableref,
                caption),level=6)
            table = doc.add_table(t.shape[0]+1,t.shape[1]+1)
            for tcell,col in zip(table.rows[0].cells[1:],t.columns):
                tcell.text = str(col)
            for trow,rrow in zip(table.rows[1:],t.to_records()):
                for tcell,rcell in zip(trow.cells,rrow):
                    tcell.text = str(rcell)

    @staticmethod
    def sectionFromFunction(function,*args,**kwargs):
        """
        This staticmethod executes the function that is passed with the provided args and kwargs.
        The first line of the function docstring is used as the section title, the comments
        within the function body are parsed and added as the section text.
        The function should return an ordered dict of figures and tables, that are then
        attached to the section.

        >>> # Section title of example function
        ... def exampleFunction(a,b=None):
        ...     'Mock figures and tables included'
        ...     figures = (('fig1',Mock()),('fig2',Mock()))
        ...     tables = (('tab1',Mock()),('tab2',Mock()))
        ...     return figures, tables
        >>> Section.sectionFromFunction(exampleFunction,Mock(),b=Mock())
        <Section @ Section title of example function>
        """
        figures, tables = function(*args,**kwargs)
        title = inspect.getcomments(function)[1:].strip()
        text = inspect.getdoc(function)
        code = inspect.getsource(function)
        return Section(title=title,text=text,figures=figures,tables=tables,code=code)

class Report(Section):
    """
    Contains as main attribute a list of sections.
    Defines methods of outputting the sections.
    outfile should not include a final extension, as
    that is determined by the different output methods.
    """
    def __init__(self,title,intro='',conclusion='',outname='',outfile=None,author=None,addTime=True):
        import time
        super().__init__(title=title,text=intro)
        self.sections = self.subs #Report sections can be accessed by both sections and subs attribute
        self.conclusion = conclusion.strip()
        self.outfile = outfile if outfile else '{}{}{}'.format(reportsDir,time.strftime('%Y_%m_%d'),
                                                               '_'+outname if outname else '')
        self.author = author
        self.addTime = addTime

        #Fig and table ref management hidden variables
        self._fignr = count(1)
        self._tabnr = count(1)
    
    def list(self):
        """
        Get an overview of the report content list
        """
        for i in range(len(self.sections)):
            self.sections[i].list(walkTrace=(i+1,))
        
    def outputZip(self,figtype='png'):
        """
        Outputs the report in a zip container.
        Figs and tabs as pngs and excells.
        """
        from zipfile import ZipFile
        with ZipFile(self.outfile+'.zip', 'w') as zipcontainer:
            with zipcontainer.open('summary.txt',mode='w') as zipf:
                zipf.write('# {}\n\n{}\n{}'.format(
                    self.title,
                    self.p,
                    ('\n## Conclusion\n' if self.conclusion else '')+self.conclusion
                ).encode())
            c = count(1)
            for section in self.sections:
                section.sectionOutZip(zipcontainer,'s{}_{}/'.format(next(c),section.title.replace(' ','_')),
                                      figtype=figtype)

    def outputPDF(self,**kwargs):
        """
        Makes a pdf report with pylatex
        **kwargs are send to doc.generate_pdf 
        -> see pylatex.Document.generate_pdf for help
        """
        import pylatex as pl
        geometry_options = {"tmargin": "2cm", "lmargin": "2cm"}
        doc = pl.Document(geometry_options=geometry_options)
        #Following option avoids float error when to many unplaced figs or tabs
        # (to force placing floats also \clearpage can be used after a section for example)
        doc.append(pl.utils.NoEscape(r'\extrafloats{100}'))
        doc.append(pl.utils.NoEscape(r'\title{'+self.title+'}'))
        if self.addTime:
            from time import localtime, strftime
            doc.append(pl.utils.NoEscape(r'\date{'+strftime("%Y-%m-%d %H:%M:%S", localtime())+r'}'))
        else: doc.append(pl.utils.NoEscape(r'\date{\today}'))
        if self.author: doc.append(pl.utils.NoEscape(r'\author{'+self.author+'}'))
        doc.append(pl.utils.NoEscape(r'\maketitle'))

        # Append introduction
        if self.p:
            with doc.create(pl.Section('Introduction')):
                doc.append(renewliner(self.p))

        # Sections
        c = count(1)
        for section in self.sections:
            section.sectionsPDF((next(c),),doc=doc)

        # Append conclusion
        if self.conclusion:
            with doc.create(pl.Section('Conclusion')):
                doc.append(renewliner(self.conclusion))

        # Generate pdf
        doc.generate_pdf(self.outfile,**kwargs)

    def outputWord(self):
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_heading(self.title, level=0)
        if self.addTime:
            from time import localtime, strftime
            doc.add_heading(strftime("%Y-%m-%d %H:%M:%S", localtime()), level=1)
            
         # Append introduction
        if self.p:
            doc.add_heading('Introduction',level=1)
            for p in renewliner(self.p).split('\n'):
                doc.add_paragraph(p)

        # Sections
        c = count(1)
        #Prepare fig and table numbers
        self.listFigures(tuple())
        self.listTables(tuple())
        for section in self.sections:
            section.sectionsWord((next(c),),doc=doc)

        # Append conclusion
        if self.conclusion:
            doc.add_heading('Conclusion', level=1)
            for p in renewliner(self.conclusion).split('\n'):
                doc.add_paragraph(p)

        # Generate Word document
        doc.save(self.outfile+'.docx')

    @staticmethod
    def getReportTable(reportzipfile,tablefilename,inReportsDir=True):
        import zipfile, io
        
        if inReportsDir: reportzipfile = reportsDir+reportzipfile
        with zipfile.ZipFile(reportzipfile) as z:
            with z.open(tablefilename) as f:
                ft = io.TextIOWrapper(f)
                return pd.read_csv(ft,index_col=0,sep=csvsep,decimal=csvdec)
                
    
# Utilities
def makeFigFromFile(filename,*args,**kwargs):
    """
    Renders an image in a matplotlib figure, so it can be added to reports 
    args and kwargs are passed to plt.subplots
    """
    img = plt.imread(filename)
    fig,ax = plt.subplots(*args,**kwargs)
    ax.axis('off')
    ax.imshow(img)
    return fig

def pdSeriesToFrame(pdseries,colname='value'):
    "Returns a series as a pd dataframe"
    return pd.DataFrame(pdseries,columns=[colname])

def renewliner(text):
    newline = re.compile(r'(\w)\n(\w)')
    return newline.subn(r'\g<1> \g<2>',text)[0]

class FigureDict(OrderedDict):
    def __init__(self, *args, section=None, **kwds):
        super().__init__(*args, **kwds)
        self._section = section 
        
    def __setitem__(self, key, figure, **kwargs):
        super().__setitem__(key,figure,**kwargs)
