#!/usr/bin/env python
# -*- coding: utf-8 -*-

import redcap, os
from docx import Document
import argparse


def redcap_project_access(API_KEY):
    """
    Gets access to redcap form
    :param API_KEY:String obtained as argparse argument object
    :return: redcap Project object
    """
    project = redcap.Project('https://redcap.vanderbilt.edu/api/', API_KEY)
    return project

def participant_filter(project,OPTIONS):
    """
    Filter participant data from export data. 
    :param data: json object containing list of  dicts
    :param participant_id: parsed from user
    no return
    """
    data = redcap_field_fetch(project)
    participant_id = splitter(OPTIONS.participant_id)
    for participant_data in data:
        for participant in participant_id:
            if(participant in participant_data.values()):
                document = document_template_open(OPTIONS.file) 
                doc = replace_field(participant_data,document,participant,OPTIONS)

def replace_field(data,document,participant,OPTIONS):
    """
    Loops through the document and replaces the field string with
    appropriate values from redcap form
    :param project: redcap.Project Object
    :param field: string object to search for in template document
    :param replacement: string of redcap field variable name
    :return: docx.Document object
    """
    replacements = splitter(OPTIONS.replacement)
    for replacement in replacements:
        for para in document.paragraphs:
            if "sex" in replacement:                #replace his/her with appropriate pronoun
                if('his/her' in para.text):
                    inline = para.runs
                    for i in range(len(inline)):
                        if( 'his/her' in inline[i].text):
                            text = inline[i].text.replace('his/her', 'her') \
                                if data.get(replacement)=='0' \
                                else inline[i].text.replace('his/her', 'his')
                            inline[i].text=text 
            else:
                if replacement in para.text:        #replace field
                    inline = para.runs
                    for i in range(len(inline)):
                        if replacement in inline[i].text:
                            text = inline[i].text.replace(replacement, data.get(replacement))
                            inline[i].text = text
        doc = replace_in_tables(data, document, replacement)
        save_doc(doc,OPTIONS.destination,participant)
    return document

def splitter(option):
    """
    Splits participant id's by ','
    :param participant_id: args parsed from user
    :return: list containing participants
    """
    if not option:
        return None 
    else:
        return option.split(',')

def save_doc(document,destination,participant_id):
    """
    Saves document with replaced fields with participant id as filename
    :param document: docx.document object
    :param destination: args parsed from user
    :param participant_id: participant id of user
    no return
    """
    filename = destination+'/'+participant_id+'.docx'
    document.save(filename)

def replace_in_tables(data,document,replacement):
    """
    Loops through tables in the document and replaces the field string with
    the replacement data from redcap form
    :param project: redcap.Project Object
    :param field: string object to search for in template document
    :param replacement: string of redcap field variable name
    :return: docx.Document object
    """
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = para.text.replace(replacement,data.get(replacement)) if not data.get(replacement)==None \
                                            else para.text.replace(replacement,'NA')
    return document

def document_template_open(file_path):
    """
    Opens the document template used for document generation
    :param doc_name: filename of template document
    :param doc_path: to template document
    :return: True if file exists
    """
    if os.path.exists(file_path):
        document = Document(file_path)
        return document
    else:
        print ("""Template file doesn't exist: Check filepath and filename provided
        Hint:Also check that suffix of filename is provided eg:docx in filename.docx""")

def redcap_field_fetch(project):
    """
    Exports data from the redcap form associated with the API Key. 
    Only the fields mentioned are pulled.
    :param project: redcap.Project object
    :param fields: fields from the redcap form that are
                   required in the document. By default, it is the global variable 
                   where the fields are mentioned. In case of new document, just change the 
                   fields accordingly
    :return:json object with a list of dicts. Each participant info is a dict
    """
    data = project.export_records(format='json')
    return data

def add_to_parser():
    """
    Method to add arguments to default parser for Document generator.

    :return: parser object with new arguments
    """
    parser = argparse.ArgumentParser(prog='DecLearn', description="Document Generator")
    parser.add_argument("-k","--key", dest='redcap_key', default=None, help='Redcap API Token')
    parser.add_argument("-r","--replace", dest='replacement',default=None, help='Combination of  \
                        words to be replaced followed by redcap field name')
    parser.add_argument("-s", "--same", dest ='same',default=None, action="store_true",\
                        help='Bool for if replacement word and redcap field have same string')
    parser.add_argument("-f","--file", dest='file',default=None, help='Path to template document')
    parser.add_argument("-p","--participant", dest='participant_id', default=None,\
                        help='Participant Id')
    parser.add_argument("-d","--dest", dest='destination', default=None, \
                        help='Destination for output files')
    return parser

def run_document_generator():
    """
    Execution point of Document Generator

    :return: None
    """
    parser = add_to_parser()
    OPTIONS = parser.parse_args()
    project = redcap_project_access(OPTIONS.redcap_key)
    participant_filter(project,OPTIONS)

if __name__ == '__main__':
    run_document_generator()
